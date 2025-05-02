from flask import Flask, render_template, request, send_from_directory
from docx import Document
from datetime import datetime, timedelta
import os
import uuid

app = Flask(__name__)

# Nastavení cest
UPLOAD_FOLDER = 'static/contracts'
SABLONY_FOLDER = 'templates_word'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Automaticky smaže soubory starší než X dní
def smazat_stare_smlouvy(cesta, max_stari_dni=7):
    threshold = datetime.now() - timedelta(days=max_stari_dni)
    for filename in os.listdir(cesta):
        filepath = os.path.join(cesta, filename)
        if os.path.isfile(filepath):
            cas_zmeny = datetime.fromtimestamp(os.path.getmtime(filepath))
            if cas_zmeny < threshold:
                os.remove(filepath)

# Nahrazení v runech se zachováním formátu
def nahrad_v_paragrafech(paragraphs, nahrady):
    for p in paragraphs:
        for run in p.runs:
            for klic, hodnota in nahrady.items():
                if klic in run.text:
                    run.text = run.text.replace(klic, hodnota)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Vymazání starých smluv
        smazat_stare_smlouvy(UPLOAD_FOLDER)

        # Získání dat z formuláře
        nazev_akce = request.form['nazev_akce']
        cislo_akce = request.form['cislo_akce']
        cislo_smlouvy = request.form['cislo_smlouvy']
        objednatel = request.form['objednatel']
        tds = request.form['tds']
        datum_input = request.form['datum']
        sablona = request.form['sablona']  # např. "smlouva_o_dilo"

        # Datum ve formátu dd. mm. rrrr
        datum = datetime.strptime(datum_input, '%Y-%m-%d').strftime('%d. %m. %Y')

        # Seznam zástupných hodnot
        nahrady = {
            '{{nazev}}': nazev_akce,
            '{{cislo}}': cislo_akce,
            '{{ID_smlouvy}}': cislo_smlouvy,
            '{{objednatel}}': objednatel,
            '{{TDS}}': tds,
            '{{datum}}': datum
        }

        # Výběr správné šablony
        sablona_path = os.path.join(SABLONY_FOLDER, sablona + '.docx')
        if not os.path.exists(sablona_path):
            return "Šablona neexistuje.", 400

        # Načtení a zpracování Word dokumentu
        doc = Document(sablona_path)
        nahrad_v_paragrafech(doc.paragraphs, nahrady)
        for section in doc.sections:
            nahrad_v_paragrafech(section.footer.paragraphs, nahrady)

        # Uložení souboru s unikátním názvem
        filename = f"{sablona}_{uuid.uuid4().hex}.docx"
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        doc.save(filepath)

        return render_template('form.html', download_link='/' + filepath)

    return render_template('form.html')


@app.route('/static/contracts/<filename>')
def download_file(filename):
    return send_from_directory(UPLOAD_FOLDER, filename, as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)
